import os
import sys
import subprocess
import time
import requests
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import re
import pygame
from gtts import gTTS
from io import BytesIO
import speech_recognition as sr
from datetime import datetime

# Тимчасово вимикаємо стандартний потік виводу
def suppress_pygame_logs():
    sys.stdout = open(os.devnull, "w")
    sys.stderr = open(os.devnull, "w")

# Відновлюємо стандартний потік виводу
def restore_stdout():
    sys.stdout = sys.__stdout__
    sys.stderr = sys.__stderr__

# Перед ініціалізацією pygame
suppress_pygame_logs()
restore_stdout()

class VoiceAssistant:
    def __init__(self):
        self.is_speaking = False
        self.recognizer = sr.Recognizer()

    # Метод для голосового відтворення
    def speak(self, text):
        if self.is_speaking:
            return

        self.is_speaking = True
        try:
            tts = gTTS(text=text, lang='uk')
            audio = BytesIO()
            tts.write_to_fp(audio)
            audio.seek(0)

            pygame.mixer.init()
            pygame.mixer.music.load(audio, "mp3")
            pygame.mixer.music.play()

            while pygame.mixer.music.get_busy():
                pygame.time.wait(100)

            audio.close()
        finally:
            self.is_speaking = False
            pygame.mixer.quit()

    # Метод для прослуховування команд
    def listen_command(self):
        try:
            with sr.Microphone() as mic:
                self.recognizer.adjust_for_ambient_noise(mic, duration=0.5)
                print("Очікую команду...")
                audio = self.recognizer.listen(mic, timeout=10)
                command = self.recognizer.recognize_google(audio, language='uk-UA').lower()
                print(f"Ви сказали: {command}")
                return command
        except sr.UnknownValueError:
            return ''
        except Exception as e:
            print(f"Помилка: {e}")
            return ''

    # Функція для очищення назв файлів
    def sanitize_filename(self, filename):
        return re.sub(r'[\\/*?:"<>|]', "", filename)

    # Команда: Запис аудіо
    def record_audio(self):
        self.speak("Скажіть, як назвати аудіофайл.")
        filename = self.listen_command()
        filename = self.sanitize_filename(filename) or "recorded_audio"
        file_path = os.path.join(os.path.expanduser("~"), 'Downloads', f"{filename}.wav")
        self.speak("Записую аудіо, скажіть що потрібно записати.")
        with sr.Microphone() as mic:
            audio = self.recognizer.listen(mic)
            with open(file_path, 'wb') as f:
                f.write(audio.get_wav_data())
        self.speak("Аудіо записано і збережено у ваші завантаження.")

    # Команда: Запис тексту
    def record_text(self):
        self.speak("Скажіть, як назвати текстовий файл.")
        filename = self.listen_command()
        filename = self.sanitize_filename(filename) or "recorded_text"
        file_path = os.path.join(os.path.expanduser("~"), 'Downloads', f"{filename}.docx")
        self.speak("Що потрібно записати текстом?")
        text = self.listen_command().capitalize() + "."
        document = Document()
        document.add_paragraph(text)
        document.save(file_path)
        self.speak("Текст записано та збережено у ваші завантаження.")

    # Команда: Додати завдання
    def add_task(self):
        self.speak("Скажіть, як назвати документ для to-do списку.")
        filename = self.listen_command()
        filename = self.sanitize_filename(filename) or "todo_list"
        file_path = os.path.join(os.path.expanduser("~"), 'Downloads', f"{filename}.docx")

        self.speak("Що додати до списку завдань?")
        task = self.listen_command().capitalize()

        if os.path.exists(file_path):
            document = Document(file_path)
        else:
            document = Document()

        document.add_paragraph(f"• {task}!")
        document.save(file_path)

        self.speak(f'Завдання додано.')

    # Команда: Пошук у Google
    def search_google(self):
        try:
            self.speak("Що шукати в Google?")
            query = self.listen_command()
            chrome_options = Options()
            chrome_options.add_experimental_option("detach", True)
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(f"https://www.google.com/search?q={query}")
        except Exception as e:
            self.speak("Перевірте підключення до інтернету.")

    # Команда: Пошук на YouTube
    def search_youtube(self):
        try:
            self.speak("Що шукати на YouTube?")
            query = self.listen_command()
            chrome_options = Options()
            chrome_options.add_experimental_option("detach", True)
            driver = webdriver.Chrome(options=chrome_options)
            driver.get(f"https://www.youtube.com/results?search_query={query}")
        except Exception as e:
            self.speak("Перевірте підключення до інтернету.")

    # Команда: Перезавантаження ПК
    def reboot_pc(self):
    # Надсилаємо запит на введення хвилин
        self.speak("Введіть через скільки хвилин ви хочете перезавантажити комп'ютер.")
        print("PROMPT:REBOOT Введіть через скільки хвилин ви хочете перезавантажити комп'ютер.", flush=True)

    # Очікуємо на вхідне повідомлення (від Electron)
        try:
            minutes = int(input())  # Отримуємо хвилини
            seconds = minutes * 60
            subprocess.run(["shutdown", "/r", "/t", str(seconds)])
            self.speak(f"Перезавантаження комп'ютера через {minutes} хвилин.")
        except ValueError:
            self.speak("Введено неправильне значення. Скасовано.")

    # Команда: Вимкнення ПК
    def shutdown_pc(self):
        # Надсилаємо запит на введення хвилин
        self.speak("Введіть через скільки хвилин ви хочете вимкнути комп'ютер.")
        print("PROMPT:SHUTDOWN Введіть через скільки хвилин ви хочете вимкнути комп'ютер.", flush=True)
        try:
            minutes = int(input())
            seconds = minutes * 60
            subprocess.run(["shutdown", "/s", "/f", "/t", str(seconds)])
            self.speak(f"Комп'ютер вимкнеться через заданий час.")
        except ValueError:
            self.speak("Введено неправильне значення. Скасовано.")

    # Команда: Нагадування
    def set_reminder(self):
        self.speak("Про що нагадати?")
        reminder_text = self.listen_command()
        self.speak(f"Ви хочете, щоб я нагадала: {reminder_text}")
        # Надсилаємо запит на введення хвилин
        self.speak("Введіть час у хвилинах, через який нагадати.")
        print("PROMPT:REMINDER Введіть час у хвилинах, через який нагадати.", flush=True)
        try:
            minutes = int(input())
            time.sleep(minutes * 60)
            self.speak(f"Нагадую вам: {reminder_text}")
        except ValueError:
            self.speak("Введено неправильне значення. Скасовано.")

    # Команда: Отримання погоди
    def get_weather(self, city_name):
        api_key = "045a71c13debb78c124e2b0135897b1a"
        base_url = "http://api.openweathermap.org/data/2.5/weather"
        params = {'q': city_name, 'appid': api_key, 'units': 'metric', 'lang': 'uk'}
        try:
            response = requests.get(base_url, params=params)
            data = response.json()
            if data["cod"] != "404":
                main = data["main"]
                weather = data["weather"][0]
                temperature = main["temp"]
                feels_like = main["feels_like"]
                description = weather["description"]
                return (f"{city_name.capitalize()}, погода: {description}, "
                        f"температура {temperature} °C, відчувається як {feels_like}.")
            else:
                return "Місто не знайдено, будь ласка, повторіть назву міста."
        except Exception as e:
            return "Не вдалося отримати дані про погоду. Перевірте підключення до інтернету."

    # Команда: Скасувати вимкнення
    def cancel_shutdown(self):
        try:
            os.system("shutdown /a")
            self.speak("Заплановане вимкнення комп'ютера скасовано.")
        except Exception as e:
            return "Не вдалося скасувати операцію. Можливо, не було заплановано операції вимкнення."

    def cancel_reboot(self):
        try:
            os.system("shutdown /a")
            self.speak("Заплановане перезавантаження комп'ютера скасовано.")
        except Exception as e:
            return "Не вдалося скасувати операцію. Можливо, не було заплановано операції перезавантаження."

    # Команда: Час
    def get_current_time(self):
        now = datetime.now().strftime("%H:%M")
        self.speak(f"Зараз {now}")

    # Команда: Дата
    def get_current_date(self):
        now = datetime.now().strftime("%d-%m-%Y")
        self.speak(f"Сьогодні {now}")

    # Команда: Таймер
    def set_timer(self):
        self.speak("Введіть на скільки хвилин ви хочете встановити таймер?")
        # Надсилаємо запит на введення хвилин
        print("PROMPT:TIMER Введіть на скільки хвилин ви хочете встановити таймер?", flush=True)
        # Очікуємо на вхідне повідомлення (від Electron)
        try:
            minutes = int(input())
            seconds = minutes * 60
            self.speak(f"Таймер на {minutes} хвилин розпочато.")
            time.sleep(seconds)
            self.speak("Час вийшов! Таймер завершено.")
        except ValueError:
            self.speak("Будь ласка, введіть правильне число для таймера.")

    # Команда: Калькулятор
    def calculate(self):
        self.speak("Скажіть математичне рівняння для обчислення.")
        equation = self.listen_command()
        try:
            result = eval(equation)  # Обчислюємо рівняння
            self.speak(f"Результат: {result}")
        except Exception as e:
            self.speak("Виникла помилка при обчисленні. Будь ласка, спробуйте ще раз.")
